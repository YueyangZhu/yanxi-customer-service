from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from PIL import ImageColor
_original_getrgb = ImageColor.getrgb
def _safe_getrgb(color):
    if isinstance(color, str) and len(color) == 6:
        color = chr(35) + color
    return _original_getrgb(color)
ImageColor.getrgb = _safe_getrgb
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
ASSETS.mkdir(parents=True, exist_ok=True)
OUT = ROOT / "言析智能客服_Coze核心流程设计文档.docx"

NAVY = "17324D"
BLUE = "2E74B5"
CYAN = "18A7A0"
PALE = "E8EEF5"
LIGHT = "F4F6F9"
GOLD = "D89B2B"
RED = "C74747"
INK = "263645"
MUTED = "6B7785"
WHITE = "FFFFFF"
GREEN = "3D8B66"

FONT_PATH = r"C:\Windows\Fonts\msyh.ttc"
FONT_BOLD_PATH = r"C:\Windows\Fonts\msyhbd.ttc"


def rgb(hexstr):
    return tuple(int(hexstr[i:i+2], 16) for i in (0, 2, 4))


def pil_font(size, bold=False):
    path = FONT_BOLD_PATH if bold and Path(FONT_BOLD_PATH).exists() else FONT_PATH
    return ImageFont.truetype(path, size)


def rounded(draw, box, fill, outline=None, radius=18, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=rgb(fill), outline=rgb(outline) if outline else None, width=width)


def center_text(draw, box, text, size=25, color=INK, bold=False):
    font = pil_font(size, bold)
    lines = text.split("\n")
    heights = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=font)
        heights.append(b[3] - b[1])
    total = sum(heights) + (len(lines)-1)*8
    y = box[1] + (box[3]-box[1]-total)/2
    for line, h in zip(lines, heights):
        b = draw.textbbox((0, 0), line, font=font)
        w = b[2]-b[0]
        draw.text((box[0]+(box[2]-box[0]-w)/2, y), line, font=font, fill=rgb(color))
        y += h + 8


def arrow(draw, a, b, color=BLUE, width=5):
    draw.line([a, b], fill=rgb(color), width=width)
    import math
    ang = math.atan2(b[1]-a[1], b[0]-a[0])
    l = 16
    p1 = (b[0]-l*math.cos(ang-0.55), b[1]-l*math.sin(ang-0.55))
    p2 = (b[0]-l*math.cos(ang+0.55), b[1]-l*math.sin(ang+0.55))
    draw.polygon([b, p1, p2], fill=rgb(color))


def save_overall():
    im = Image.new("RGB", (1500, 700), "FFFFFF")
    d = ImageDraw.Draw(im)
    nodes = [
        ((60, 90, 300, 210), "用户网页", BLUE),
        ((390, 90, 660, 210), "CloudBase\n云函数", NAVY),
        ((750, 90, 1030, 210), "Coze\n主对话流", CYAN),
        ((1120, 90, 1430, 210), "知识库 /\n子工作流", GREEN),
        ((390, 400, 660, 520), "CloudBase\n数据库", NAVY),
        ((750, 400, 1030, 520), "人工客服\n工作台", GOLD),
        ((1120, 400, 1430, 520), "飞书重点\n数据同步", RED),
    ]
    for box, label, col in nodes:
        rounded(d, box, LIGHT, col, 20, 4)
        center_text(d, box, label, 30, col, True)
    arrow(d, (300,150), (390,150)); arrow(d,(660,150),(750,150)); arrow(d,(1030,150),(1120,150))
    arrow(d,(525,210),(525,400)); arrow(d,(660,460),(750,460)); arrow(d,(1030,460),(1120,460))
    arrow(d,(885,400),(885,220), GOLD)
    d.text((60, 610), "主数据源：CloudBase｜AI 理解与回答：Coze｜重点协作与运营提醒：飞书", font=pil_font(27, True), fill=rgb(NAVY))
    p = ASSETS / "01_overall_architecture.png"; im.save(p); return p


def save_resource_map():
    im = Image.new("RGB", (1500, 760), "FFFFFF")
    d = ImageDraw.Draw(im)
    center = (530, 270, 970, 440)
    rounded(d, center, NAVY, NAVY, 28, 3); center_text(d, center, "after_sales_chat\n主对话流", 34, WHITE, True)
    items = [
        ((60,70,410,200), "understand_message\n消息理解", CYAN),
        ((1090,70,1440,200), "售后知识库\n政策与 FAQ", BLUE),
        ((60,540,410,680), "generate_handoff_summary\n转人工摘要", GOLD),
        ((1090,540,1440,680), "analyze_service_result\n服务结果分析", GREEN),
    ]
    for box, label, col in items:
        rounded(d, box, LIGHT, col, 22, 4); center_text(d, box, label, 27, col, True)
    arrow(d,(410,135),(530,300),CYAN); arrow(d,(970,300),(1090,135),BLUE)
    arrow(d,(530,405),(410,610),GOLD); arrow(d,(970,405),(1090,610),GREEN)
    d.text((415, 705), "入口编排 → 意图识别 → 分支执行 → 人工兜底 → 结果复盘", font=pil_font(28, True), fill=rgb(INK))
    p = ASSETS / "02_coze_resource_map.png"; im.save(p); return p


def save_chat_flow():
    im = Image.new("RGB", (1500, 1000), "FFFFFF")
    d = ImageDraw.Draw(im)
    top = [((60,40,300,140),"开始",BLUE),((380,40,760,140),"调用消息理解",CYAN),((840,40,1210,140),"按意图路由",NAVY)]
    for box,label,col in top:
        rounded(d,box,LIGHT,col,20,4); center_text(d,box,label,28,col,True)
    arrow(d,(300,90),(380,90)); arrow(d,(760,90),(840,90))
    branches = [
      ((80,260,390,370),"需要转人工\nneed_handoff = true",RED),
      ((440,260,750,370),"政策咨询\nproduct_consultation",BLUE),
      ((800,260,1110,370),"订单 / 退款\n3 类业务意图",CYAN),
      ((1160,260,1440,370),"其他问题\nelse",MUTED),
    ]
    for box,label,col in branches:
        rounded(d,box,LIGHT,col,18,3); center_text(d,box,label,23,col,True)
        arrow(d,(1025,140),((box[0]+box[2])//2,260),col,4)
    downstream = [
      ((80,490,390,600),"回复转人工提示",RED),
      ((440,450,750,540),"检索售后知识",BLUE),
      ((440,600,750,690),"生成知识回答",BLUE),
      ((800,520,1110,630),"生成业务引导",CYAN),
      ((1160,490,1440,600),"回复转人工提示",RED),
    ]
    for box,label,col in downstream:
        rounded(d,box,LIGHT,col,18,3); center_text(d,box,label,23,col,True)
    arrow(d,(235,370),(235,490),RED); arrow(d,(595,370),(595,450),BLUE); arrow(d,(595,540),(595,600),BLUE)
    arrow(d,(955,370),(955,520),CYAN); arrow(d,(1300,370),(1300,490),RED)
    agg=(520,800,980,920); rounded(d,agg,NAVY,NAVY,24,3); center_text(d,agg,"汇总最终回复 → 结束",32,WHITE,True)
    for x,y in [(235,600),(595,690),(955,630),(1300,600)]: arrow(d,(x,y),(750,800),NAVY,4)
    p=ASSETS/"03_after_sales_chat.png"; im.save(p); return p


def save_sequence():
    im=Image.new("RGB",(1500,880),"FFFFFF"); d=ImageDraw.Draw(im)
    names=[("用户",150),("网页/云函数",480),("Coze",800),("CloudBase",1130),("人工客服",1390)]
    for n,x in names:
        d.text((x-55,30),n,font=pil_font(24,True),fill=rgb(NAVY)); d.line((x,80,x,820),fill=rgb(PALE),width=4)
    events=[
      (130, "发送问题",150,480,BLUE),(220,"调用对话流",480,800,CYAN),(310,"返回 answer + action",800,480,CYAN),
      (400,"保存消息 / 查询订单",480,1130,NAVY),(490,"低置信度或投诉：创建工单",480,1130,RED),
      (580,"客服接单并回复",1390,1130,GOLD),(670,"用户查看人工回复",1130,150,GOLD),(760,"评价 → 服务结果分析",150,800,GREEN)
    ]
    for y,label,x1,x2,col in events:
        arrow(d,(x1,y),(x2,y),col,4)
        mid=(x1+x2)//2; b=d.textbbox((0,0),label,font=pil_font(21,True)); w=b[2]-b[0]
        d.rectangle((mid-w/2-10,y-37,mid+w/2+10,y-8),fill="FFFFFF")
        d.text((mid-w/2,y-35),label,font=pil_font(21,True),fill=rgb(col))
    p=ASSETS/"04_end_to_end_sequence.png"; im.save(p); return p


def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); mar=tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar'); tcPr.append(mar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=mar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); mar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')


def set_col_width(cell, dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(dxa)); tcW.set(qn('w:type'),'dxa')


def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)


def set_run(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name='Microsoft YaHei'; run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
    run.font.size=Pt(size); run.font.color.rgb=RGBColor(*rgb(color)); run.bold=bold; run.italic=italic


def add_p(doc, text='', size=10.5, color=INK, bold=False, after=6, align=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.25
    if align is not None: p.alignment=align
    set_run(p.add_run(text),size,color,bold); return p


_CALLOUT_COUNT = 0
def add_callout(doc, title, body, color=BLUE):
    global _CALLOUT_COUNT
    _CALLOUT_COUNT += 1
    if _CALLOUT_COUNT == 1 or _CALLOUT_COUNT == 6:
        return
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    set_col_width(t.cell(0,0),9360); shade(t.cell(0,0),LIGHT); set_cell_margins(t.cell(0,0),140,180,140,180)
    p=t.cell(0,0).paragraphs[0]; p.paragraph_format.space_after=Pt(4)
    set_run(p.add_run(title+'  '),11,color,True); set_run(p.add_run(body),10.3,INK)
    add_p(doc,'',after=2)


def add_table(doc, headers, rows, widths=None, font=8.5):
    table=doc.add_table(rows=1,cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    widths=widths or [9360//len(headers)]*len(headers)
    set_repeat_table_header(table.rows[0])
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i]; set_col_width(c,widths[i]); shade(c,PALE); set_cell_margins(c)
        c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
        set_run(p.add_run(str(h)),font+0.3,NAVY,True)
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row):
            set_col_width(cells[i],widths[i]); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.1
            set_run(p.add_run(str(val)),font,INK)
    add_p(doc,'',after=2); return table


def heading(doc,text,level=1):
    return doc.add_heading(text,level=level)


def add_image(doc,path,width=6.45):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(6)
    p.add_run().add_picture(str(path),width=Inches(width))


_PAGE_BREAK_COUNT = 0
def page_break(doc):
    global _PAGE_BREAK_COUNT
    if _PAGE_BREAK_COUNT == 0:
        doc.add_page_break()
    _PAGE_BREAK_COUNT += 1


def set_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=paragraph.add_run('第 '); set_run(r,8.5,MUTED)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); paragraph._p.append(fld)
    r=paragraph.add_run(' 页'); set_run(r,8.5,MUTED)


def build():
    overall, resource, chat, sequence = save_overall(), save_resource_map(), save_chat_flow(), save_sequence()
    doc=Document(); sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.492)
    normal=doc.styles['Normal']; normal.font.name='Microsoft YaHei'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
    normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor(*rgb(INK)); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name,size,col,before,after in [('Title',30,NAVY,0,8),('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,NAVY,10,5)]:
        s=doc.styles[name]; s.font.name='Microsoft YaHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor(*rgb(col)); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    header=sec.header.paragraphs[0]; set_run(header.add_run('言析智能客服｜Coze 核心流程设计'),8.5,MUTED,True)
    set_page_number(sec.footer.paragraphs[0])

    add_p(doc,'AI 产品经理作品集 · 系统设计文档',11,GOLD,True,18)
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.add_run('言析电商售后智能客服')
    add_p(doc,'Coze 核心工作流、对话流与变量字典',16,NAVY,True,12)
    add_p(doc,'从“用户提问”到“AI 回答、人工兜底、服务评价与知识回流”的完整设计说明',11.5,MUTED,False,26)
    add_image(doc,resource,6.3)
    add_table(doc,['文档项','内容'],[
      ('文档版本','v1.0'),('项目定位','AI 产品经理面试作品集 · 电商售后客服'),('当前 Coze 版本','understand_message v0.1.1；after_sales_chat v0.1.0；两个分析工作流 v0.1.0'),('更新日期','2026-06-20')],[2700,6660],9.2)
    add_callout(doc,'一句话理解','对话流负责“接住并编排一轮对话”，工作流负责“把某项能力做成稳定、可复用、结构化的模块”。',CYAN)

    page_break(doc); heading(doc,'1. 文档阅读地图',1)
    add_p(doc,'这份文档既是搭建说明，也是面试讲解底稿。建议先看业务闭环，再按主对话流和三个工作流理解节点关系。')
    add_table(doc,['章节','解决的问题','重点产物'],[
      ('2 总体设计','系统如何形成业务闭环？','架构图、职责边界'),('3 Coze 资源关系','为什么既有对话流又有工作流？','资源关系图'),('4 消息理解','如何识别意图、订单号、风险？','结构化识别结果'),('5 主对话流','如何路由知识、订单、退款与人工？','分支编排'),('6 转人工摘要','如何让人工快速接手？','工单摘要字段'),('7 服务结果分析','如何从评价形成优化闭环？','失败原因与知识缺口'),('8 变量字典','英文字段分别是什么意思？','中英文术语表'),('9 端到端示例','真实交互如何流转？','时序图与案例')],[1300,4200,3860],8.8)
    heading(doc,'2. 总体业务闭环与技术分工',1)
    add_image(doc,overall,6.45)
    add_table(doc,['系统模块','核心职责','不负责的事情'],[
      ('Coze','意图/实体/情绪/风险识别；知识问答；对话路由；人工摘要与结果分析','不保存完整业务主数据，不直接承诺订单或退款结果'),('CloudBase','网页 API、会话与消息、模拟订单、退款、工单、评价、数据分析','不承担自然语言理解和知识生成'),('飞书','同步高优工单、低评分、知识缺口、运营日报','不是完整消息数据库，不作为唯一数据源'),('React 网页','用户聊天、退款表单、工单状态、客服工作台、运营看板','不保存 Coze 密钥，不直接调用敏感服务')],[1250,4150,3960],8.4)

    page_break(doc); heading(doc,'3. Coze 核心资源与关联逻辑',1)
    add_image(doc,resource,6.45)
    add_table(doc,['资源','类型','什么时候调用','输出被谁使用'],[
      ('after_sales_chat','对话流 Dialogue Flow','用户每发送一条消息','网页最终展示 answer，并按 action 执行业务动作'),('understand_message','工作流 Workflow','主对话流处理每条用户消息时','按意图路由节点'),('言析电商售后知识库','Knowledge Base','product_consultation 分支','生成知识回答节点'),('generate_handoff_summary','工作流 Workflow','需要创建人工工单时','CloudBase 工单与人工工作台'),('analyze_service_result','工作流 Workflow','用户评价或会话结束后','数据看板、知识缺口与优化任务')],[1900,1400,3000,3060],8.2)
    add_callout(doc,'设计原则','识别、生成、业务查询和数据存储解耦。模型只判断与表达；订单、退款、工单等事实由业务系统处理。',NAVY)
    heading(doc,'3.1 为什么智能客服不只使用对话流？',2)
    add_table(doc,['对话流适合','工作流适合'],[
      ('处理一轮对话、控制分支、返回用户可见消息','执行可复用能力、固定输入输出、独立测试与版本发布'),('记住对话上下文、汇总不同分支回答','消息理解、转人工摘要、服务结果分析'),('面向用户体验','面向系统稳定性与复用性')],[4680,4680],9.0)

    page_break(doc); heading(doc,'4. 工作流一：understand_message（消息理解）',1)
    add_p(doc,'目标：把自然语言问题转换成稳定的结构化数据，供后续分支和业务系统直接使用。')
    add_table(doc,['节点顺序','节点名称','作用'],[
      ('1','开始 Start','接收 message 与 history_summary'),('2','识别意图与风险 LLM','识别意图、订单号、情绪、风险、置信度和下一步动作'),('3','解析识别结果 JavaScript','把模型 JSON 字符串转换成强类型字段'),('4','结束 End','返回 8 个标准字段')],[1100,2500,5760],9.0)
    heading(doc,'4.1 输入参数',2)
    add_table(doc,['字段','中文含义','类型','必填','来源'],[
      ('message','用户当前消息','String','是','主对话流 USER_INPUT'),('history_summary','必要的历史摘要','String','否','主对话流或云函数生成')],[1900,2600,1300,1000,2560],8.8)
    heading(doc,'4.2 输出参数',2)
    add_table(doc,['字段','中文含义','示例','用途'],[
      ('intent','识别意图','logistics_query','决定进入哪个分支'),('confidence','置信度（0–1）','0.90','低于阈值时转人工'),('order_no','订单号','OD20260620001','传给订单/退款查询'),('emotion','用户情绪','anxious','辅助判断服务风险'),('risk_level','风险等级','high','决定优先级和是否转人工'),('need_handoff','是否需要人工','true','最高优先级分支条件'),('handoff_reason','转人工原因','用户投诉且金额有争议','写入工单'),('action','下一步动作','query_order','通知网页或业务系统执行')],[1700,2300,2300,3060],8.3)
    heading(doc,'4.3 意图与动作映射',2)
    add_table(doc,['intent','中文','action','后续处理'],[
      ('product_consultation','商品/政策咨询','search_knowledge','检索知识库'),('logistics_query','物流/订单查询','query_order','CloudBase 查询订单'),('refund_progress','退款进度','query_refund','CloudBase 查询退款'),('refund_apply','申请退款','show_refund_form','网页展示退款表单'),('complaint','投诉/强烈不满','create_ticket','创建人工工单'),('human_service','主动要求人工','create_ticket','创建人工工单'),('unknown','无法识别','create_ticket','记录知识缺口并转人工')],[2350,1800,2200,3010],8.2)

    page_break(doc); heading(doc,'5. 主对话流：after_sales_chat',1)
    add_p(doc,'主对话流是用户入口。它不直接完成所有能力，而是调用消息理解工作流后，根据结构化结果选择最合适的处理分支。')
    add_image(doc,chat,6.3)
    heading(doc,'5.1 分支优先级（从上到下匹配）',2)
    add_table(doc,['优先级','条件','处理链路','最终回答'],[
      ('1','need_handoff = true','回复转人工提示','告知已转人工并创建工单'),('2','intent = product_consultation','检索售后知识 → 生成知识回答 → 回复知识答案','基于知识库的政策说明'),('3','intent = logistics_query','生成业务引导 → 回复业务引导','确认订单号并提示查询'),('4','intent = refund_progress','生成业务引导 → 回复业务引导','确认订单号并提示查询退款'),('5','intent = refund_apply','生成业务引导 → 回复业务引导','引导网页展示退款表单'),('Else','其他情况','回复转人工提示','兜底，避免无答案')],[900,2500,3350,2610],8.0)
    add_callout(doc,'为什么 need_handoff 放第一位？','投诉、金额争议、主动要求人工等高风险信号必须覆盖普通意图，避免系统把投诉误送入普通知识问答。',RED)
    heading(doc,'5.2 知识问答分支',2)
    add_table(doc,['节点','关键配置','说明'],[
      ('检索售后知识','混合检索；最多 5 条；最低匹配度 0.50；查询改写/结果重排开启','提高召回率并减少无关切片'),('生成知识回答','只依据 knowledge_results；禁止编造订单、金额、物流和退款结果','知识不足时明确说明并转人工'),('回复知识答案','role = assistant；content = LLM.output','把最终文本写入当前对话')],[1750,3950,3660],8.3)
    heading(doc,'5.3 最终回答汇总',2)
    add_p(doc,'变量聚合节点“汇总最终回复”依次读取：转人工提示 content、知识答案 content、业务引导 content，并返回每个分组中第一个非空值。结束节点把 Group1 作为 output 返回。')

    page_break(doc); heading(doc,'6. 工作流二：generate_handoff_summary（转人工摘要）',1)
    add_p(doc,'目标：把完整对话压缩成客服可在 10 秒内读完的交接摘要，避免用户重复描述。')
    add_table(doc,['输入字段','中文含义','是否必填'],[
      ('conversation_history','完整会话文本','是'),('intent','识别意图','是'),('order_no','订单号','否'),('emotion','情绪','是'),('risk_level','风险等级','是'),('handoff_reason','转人工原因','是')],[2800,4660,1900],9.0)
    add_table(doc,['输出字段','中文含义','进入工单后的用途'],[
      ('summary','对话摘要','工单摘要区'),('customer_request','客户最终诉求','客服首先确认处理目标'),('key_facts','关键事实','订单号、等待时长、争议点'),('completed_actions','已完成动作','避免重复执行'),('pending_action','下一步待处理事项','客服处理清单'),('priority','优先级 low/medium/high','工单排序与 SLA')],[2500,2800,4060],8.5)
    add_callout(doc,'防幻觉约束','摘要只能使用输入中明确出现的信息；不能把“准备查询”写成“已经查询完成”，不能编造订单、金额或处理结论。',GOLD)
    heading(doc,'6.1 当前已验证案例',2)
    add_p(doc,'用户反馈退款等待多日且金额不对，明确要求投诉。流程输出 priority = high、customer_request = 处理退款未到账与金额不符并接受投诉、pending_action = 核验退款审核与到账状态。')

    page_break(doc); heading(doc,'7. 工作流三：analyze_service_result（服务结果分析）',1)
    add_p(doc,'目标：在用户评价或会话结束后，判断是否解决、失败发生在哪个阶段、是否存在知识缺口，并给出优化建议。')
    add_table(doc,['输入字段','中文含义','类型'],[
      ('conversation_history','完整会话','String'),('resolved','用户是否确认解决','Boolean'),('score','服务评分 1–5','Number'),('user_comment','评价内容','String，可选'),('final_intent','最终意图','String'),('handoff_used','是否使用人工客服','Boolean')],[2700,4300,2360],8.7)
    add_table(doc,['输出字段','中文含义','可能值/示例'],[
      ('outcome','服务结果','solved / unsolved / partial'),('failure_reason','失败原因','knowledge_gap / wrong_intent / handoff_delay 等'),('responsible_stage','主要责任环节','knowledge_retrieval / human_service 等'),('knowledge_gap','是否知识缺口','true / false'),('gap_question','需要补充知识库的问题','会员生日礼物怎么领取'),('optimization','优化建议','补充知识、调整路由、改善客服 SLA'),('analysis_summary','分析摘要','用于运营后台与日报')],[2400,2800,4160],8.3)
    heading(doc,'7.1 失败原因枚举',2)
    add_table(doc,['值','中文','典型情形'],[
      ('none','无失败','已解决且评分正常'),('knowledge_gap','知识缺口','知识库没有可靠答案'),('wrong_intent','意图误判','问题进入了错误分支'),('business_system','业务系统失败','订单/退款接口异常'),('handoff_delay','人工接入延迟','已转人工但长时间无答复'),('response_quality','回答质量问题','内容不清晰、不完整或语气不当'),('policy_limit','政策限制','超出支持范围或政策不能满足')],[2200,1800,5360],8.2)
    heading(doc,'7.2 责任阶段枚举',2)
    add_table(doc,['值','中文'],[
      ('none','无责任阶段'),('intent_recognition','意图识别'),('knowledge_retrieval','知识检索'),('business_query','业务查询'),('human_service','人工客服')],[3200,6160],8.8)

    page_break(doc); heading(doc,'8. 英文变量与术语字典',1)
    add_p(doc,'英文变量是系统模块之间的“统一语言”。变量名保持英文可减少接口对接歧义，文档和界面则使用中文解释。')
    add_table(doc,['英文','中文','通俗解释'],[
      ('Workflow','工作流','把一项能力拆成稳定步骤，可独立测试和复用'),('Dialogue Flow','对话流','面向用户的一轮轮对话与分支编排'),('Agent','智能体','带角色、模型、知识和技能的 AI 应用'),('Node','节点','流程中的一个处理步骤'),('Input / Output','输入 / 输出','进入节点的数据 / 节点返回的数据'),('String / Number / Boolean','字符串 / 数字 / 布尔值','文本 / 数值 / true 或 false'),('Intent','意图','用户想做什么'),('Entity','实体','订单号、商品名等关键对象'),('Confidence','置信度','模型对识别结果有多确定'),('Action','动作','网页或业务系统下一步做什么'),('Handoff','转人工','从 AI 服务切换到人工客服'),('Risk level','风险等级','问题需要多高优先级处理'),('Knowledge gap','知识缺口','知识库未覆盖且需要补充的问题'),('Resolved','是否解决','用户是否确认问题解决'),('SLA','服务等级协议','例如人工客服应在多久内响应')],[2700,2400,4260],8.2)
    heading(doc,'8.1 常用值的中文对照',2)
    add_table(doc,['字段','英文值','中文含义'],[
      ('emotion','calm / anxious / angry','平静 / 焦虑 / 生气'),('risk_level','low / medium / high','低 / 中 / 高'),('priority','low / medium / high','普通 / 较急 / 紧急'),('outcome','solved / partial / unsolved','已解决 / 部分解决 / 未解决'),('need_handoff','true / false','需要人工 / 不需要人工')],[1900,3700,3760],8.5)

    page_break(doc); heading(doc,'9. 端到端业务时序与案例',1)
    add_image(doc,sequence,6.45)
    heading(doc,'9.1 案例 A：知识咨询',2)
    add_table(doc,['步骤','系统行为'],[
      ('1','用户询问“签收后几天可以无理由退货？”'),('2','understand_message 返回 product_consultation + search_knowledge'),('3','主对话流检索售后知识库，生成基于政策的答案'),('4','CloudBase 保存用户消息、AI 回答与知识来源'),('5','用户确认是否解决并评分')],[1000,8360],8.8)
    heading(doc,'9.2 案例 B：订单物流',2)
    add_table(doc,['步骤','系统行为'],[
      ('1','用户提供订单号并询问到哪里了'),('2','识别结果为 logistics_query + query_order'),('3','Coze 只返回查询引导，不编造物流状态'),('4','网页按 action 调用 CloudBase 订单查询'),('5','网页把真实模拟订单信息渲染成订单卡片')],[1000,8360],8.8)
    heading(doc,'9.3 案例 C：投诉并要求人工',2)
    add_table(doc,['步骤','系统行为'],[
      ('1','识别 complaint、angry、high、need_handoff = true'),('2','主对话流优先进入人工分支并回复已转接'),('3','调用 generate_handoff_summary 生成摘要与 high 优先级'),('4','CloudBase 创建 ticket，人工工作台接单并回复'),('5','用户评价后调用 analyze_service_result，必要时记录知识缺口')],[1000,8360],8.8)

    page_break(doc); heading(doc,'10. 面试讲解框架与后续实施清单',1)
    add_callout(doc,'30 秒项目介绍','这是一个面向电商售后的智能客服系统。Coze 负责意图、知识问答和风险判断，CloudBase 负责订单、工单、评价与数据分析。系统不仅能回答问题，还形成了转人工、工单处理、用户评价和知识回流的完整闭环。',CYAN)
    heading(doc,'10.1 设计亮点',2)
    add_table(doc,['亮点','体现的 AI 产品能力'],[
      ('统一结构化输出','把模型能力变成可被前端和业务系统稳定消费的接口'),('高风险优先转人工','把用户体验、合规与业务风险放在普通问答之前'),('模型不编造业务事实','订单、退款结果必须来自业务系统'),('CloudBase 唯一主数据源','避免数据库与飞书重复存储导致不一致'),('服务结果自动归因','把差评转化为知识、流程或人工服务优化任务'),('可评测、可迭代','每个工作流可独立测试，主对话流可用 50+ 问题回归')],[3100,6260],8.8)
    heading(doc,'10.2 下一阶段实施顺序',2)
    add_table(doc,['顺序','任务','验收结果'],[
      ('1','发布并获取 Coze API 调用参数','云函数可获得真实 answer / action'),('2','完成 CloudBase 数据库与云函数','会话、消息、订单、退款、工单、评价可读写'),('3','接通 React 用户客服页','知识、订单、退款、人工提示均可交互'),('4','完成人工客服工作台','接单、回复、关闭、用户确认形成闭环'),('5','完成运营看板与飞书同步','展示解决率、转人工率、满意度、知识缺口'),('6','执行不少于 50 条问题评测','修正意图、知识、风险和回答问题'),('7','部署 CloudBase 并录制演示','浏览器公开链接 + 作品集演示视频')],[900,4300,4160],8.5)
    add_callout(doc,'当前状态','Coze 的主对话流、消息理解、转人工摘要、服务结果分析和售后知识库已完成第一版并通过基础测试；下一步进入 API 与 CloudBase 集成。',GREEN)

    for row in doc.tables[-1].rows:
        for cell in row.cells:
            set_cell_margins(cell, top=40, bottom=40)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        trailing = doc.paragraphs[-1]._element
        trailing.getparent().remove(trailing)
    # Keep tables together where possible and set document metadata.
    props=doc.core_properties; props.title='言析电商售后智能客服：Coze 核心流程设计文档'; props.subject='AI 产品经理作品集系统设计'; props.author='言析智能客服项目'
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__': build()
